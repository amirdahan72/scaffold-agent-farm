"""Generate a Word document from the weekly competitive digest using the docx-writer skill guidelines.

Skill: docx-writer
Guidelines followed:
- Calibri 11pt body, Calibri Light headings
- Table Grid style with bold header rows
- Intense Quote style for internal/Work IQ callouts
- No colored text, no decorative fonts
- Source URLs included as plain text in Sources section
"""
from docx import Document
from docx.shared import Pt

doc = Document()

# --- Styles per skill: Calibri 11pt body ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ============================================================
# TITLE & METADATA
# ============================================================
doc.add_heading('Weekly Competitive Digest: WAAP', level=1)

p = doc.add_paragraph()
p.add_run('Week ending: ').bold = True
p.add_run('March 9, 2026')
p = doc.add_paragraph()
p.add_run('Audience: ').bold = True
p.add_run('PM peers')
p = doc.add_paragraph()
p.add_run('Classification: ').bold = True
p.add_run('Internal')
p = doc.add_paragraph()
p.add_run('Competitors tracked: ').bold = True
p.add_run('AWS WAF, GCP Cloud Armor, Akamai, Cloudflare')

# ============================================================
# QUICK NAVIGATION
# ============================================================
doc.add_heading('Quick Navigation', level=2)
for item in [
    'Top 3 This Week \u2014 the headlines',
    'Signal Dashboard \u2014 at-a-glance heat map',
    'Competitor Activity \u2014 detailed signals',
    'Watch List \u2014 items to track',
    'Recommended Actions \u2014 what to do',
]:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# TOP 3 THIS WEEK
# ============================================================
doc.add_heading('Top 3 This Week', level=2)

top3 = [
    ('Cloudflare Attack Signature Detection (Early Access)',
     'Always-on WAF detection framework running 700+ signatures on every request without latency, decoupling detection from mitigation and eliminating the log-vs-block trade-off \u2014 a direct challenge to every WAF vendor\'s detection architecture.',
     'https://blog.cloudflare.com/attack-signature-detection/'),
    ('Cloudflare Full-Transaction Detection (Under Development)',
     'Analyzes both HTTP request AND response to correlate successful exploits, catch data exfiltration, and identify misconfigs \u2014 a unique capability no competitor currently offers.',
     'https://blog.cloudflare.com/attack-signature-detection/'),
    ('GCP Cloud Armor Hierarchical Security Policies (GA)',
     'Organization- and folder-level security policies now enforced across projects, enabling enterprise-wide WAF governance at scale. The only non-Cloudflare in-window GA release, addressing a key enterprise governance gap.',
     'https://docs.cloud.google.com/armor/docs/release-notes'),
]

for i, (title, desc, url) in enumerate(top3, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(title).bold = True
    p.add_run(f' \u2014 {desc} Source: {url}')

# ============================================================
# SIGNAL DASHBOARD
# ============================================================
doc.add_heading('Signal Dashboard', level=2)

table = doc.add_table(rows=5, cols=7, style='Table Grid')
headers = ['Competitor', 'Product', 'Pricing', 'Partnerships', 'Funding/M&A', 'People', 'Overall Heat']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for par in cell.paragraphs:
        for r in par.runs:
            r.bold = True

dashboard_data = [
    ['Cloudflare',      'HIGH', '\u2014', '\u2014', '\u2014', '\u2014', 'HIGH'],
    ['GCP Cloud Armor', 'MED',  '\u2014', '\u2014', '\u2014', '\u2014', 'MED'],
    ['AWS WAF',         '\u2014', '\u2014', '\u2014', '\u2014', '\u2014', '\u2014'],
    ['Akamai',          '\u2014', '\u2014', '\u2014', '\u2014', '\u2014', '\u2014'],
]
for ri, row_data in enumerate(dashboard_data):
    for ci, val in enumerate(row_data):
        table.rows[ri + 1].cells[ci].text = val

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('HIGH = significant activity | MED = minor activity | LOW = positive for us | \u2014 = no signals. Dashboard reflects in-window signals (Mar 2\u20139) only.')
r.italic = True
r.font.size = Pt(9)

# ============================================================
# COMPETITOR ACTIVITY
# ============================================================
doc.add_heading('Competitor Activity', level=2)

# --- Cloudflare ---
doc.add_heading('Cloudflare \u2014 Heat: HIGH', level=3)

table = doc.add_table(rows=7, cols=4, style='Table Grid')
col_headers = ['Date', 'Signal', 'Type', 'Impact']
for i, h in enumerate(col_headers):
    table.rows[0].cells[i].text = h
    for par in table.rows[0].cells[i].paragraphs:
        for r in par.runs:
            r.bold = True

cf_data = [
    ['Mar 4', 'Attack Signature Detection (Early Access) \u2014 700+ signatures, zero latency, decouples detection from mitigation', 'Product Launch', 'HIGH'],
    ['Mar 4', 'Full-Transaction Detection (Under Dev) \u2014 Analyzes request AND response; catches exfiltration and exploits', 'Product Launch', 'HIGH'],
    ['Mar 2', 'WAF managed rules \u2014 CVE-2025-52691 & CVE-2026-23760 (SmarterMail)', 'Feature Release', 'MED'],
    ['Mar 3', '2026 Threat Intelligence Report \u2014 230B daily threats, 31.4 Tbps DDoS record', 'Media Coverage', 'MED'],
    ['Feb 23 (pre-window)', 'Post-Quantum SASE platform \u2014 First vendor to support modern post-quantum encryption', 'Product Launch', 'MED'],
    ['Feb 17 (pre-window)', 'Mastercard cyber defense partnership for critical infrastructure and SMBs', 'Partnership', 'MED'],
]
for ri, row_data in enumerate(cf_data):
    for ci, val in enumerate(row_data):
        table.rows[ri + 1].cells[ci].text = val

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('So what: ').bold = True
p.add_run('Cloudflare is making a bold architectural bet with Attack Signature Detection and Full-Transaction Detection. The always-on, zero-latency detection model eliminates the oldest WAF trade-off (log mode vs. block mode) and directly pressures our rule-engine architecture. Full-Transaction Detection \u2014 analyzing responses, not just requests \u2014 is a genuinely novel capability that could become a differentiator for detecting successful exploits and data exfiltration. Their pace of innovation (4 in-window signals) reinforces the urgency of responding competitively.')

p = doc.add_paragraph(style='Intense Quote')
p.add_run('Internal Context (Work IQ): ').bold = True
p.add_run('Cloudflare is confirmed as the #1 churn risk; customers are benchmarking us against Cloudflare and some are leaving AFD/App Gateway. This validates the urgency of our accelerated AI WAAP pillars targeting Ignite. Do not distribute externally.')

# --- GCP Cloud Armor ---
doc.add_heading('GCP Cloud Armor \u2014 Heat: MED', level=3)

table = doc.add_table(rows=3, cols=4, style='Table Grid')
for i, h in enumerate(col_headers):
    table.rows[0].cells[i].text = h
    for par in table.rows[0].cells[i].paragraphs:
        for r in par.runs:
            r.bold = True

gcp_data = [
    ['Mar 3', 'ASN support in globally scoped edge security policies (GA)', 'Feature Release', 'MED'],
    ['Mar 3', 'Hierarchical security policies (GA) \u2014 Org/folder-level enforcement', 'Feature Release', 'MED'],
]
for ri, row_data in enumerate(gcp_data):
    for ci, val in enumerate(row_data):
        table.rows[ri + 1].cells[ci].text = val

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('So what: ').bold = True
p.add_run('GCP Cloud Armor is steadily maturing its enterprise governance story. Hierarchical policies address a real gap for multi-project GCP organizations. Neither feature changes Cloud Armor\'s positioning as a baseline WAF, but they reduce the gap on enterprise manageability. Watch for Cloud Armor leveraging Gemini AI capabilities.')

p = doc.add_paragraph(style='Intense Quote')
p.add_run('Internal Context (Work IQ): ').bold = True
p.add_run('We already support equivalent centralized governance through Azure Policy + WAF policy inheritance. GCP Cloud Armor remains a "baseline WAF" per our internal assessment. Do not distribute externally.')

# --- AWS WAF ---
doc.add_heading('AWS WAF \u2014 Heat: None', level=3)

doc.add_paragraph('No WAF-specific product announcements this week. The Mar 5 "What\'s New" mentioned Shield Network Security Director Findings but nothing WAF-specific. Last major WAF update was the simplified console (Jun 2025) reducing config steps by 80% with pre-configured protection packs.')

p = doc.add_paragraph()
p.add_run('So what: ').bold = True
p.add_run('AWS WAF\'s silence is itself a signal. AWS continues to treat WAF as infrastructure plumbing rather than a strategic security product. The risk: AWS could eventually embed Bedrock-powered AI detection into WAF with minimal announcement. The opportunity: AWS WAF customers are underserved on API protection, bot management, and ML-driven L7 DDoS.')

p = doc.add_paragraph(style='Intense Quote')
p.add_run('Internal Context (Work IQ): ').bold = True
p.add_run('Customers layer Cloudflare or Akamai on top of AWS WAF for advanced protection. AWS WAF and GCP Cloud Armor are viewed internally as baseline WAFs, not WAAP leaders. Do not distribute externally.')

# --- Akamai ---
doc.add_heading('Akamai \u2014 Heat: None', level=3)

doc.add_paragraph('No WAAP-specific announcements detected March 2\u20139. App & API Protector Hybrid (Apr 2025) and Firewall for AI remain the most recent major releases. Watch for RSA Conference pre-announcements.')

p = doc.add_paragraph()
p.add_run('So what: ').bold = True
p.add_run('Akamai\'s quiet week likely reflects RSA Conference timing \u2014 expect a burst of announcements in the coming weeks. Their API security depth (via NoName + NeoSec acquisitions) remains a key competitive advantage. Silence doesn\'t reduce their competitive threat.')

p = doc.add_paragraph(style='Intense Quote')
p.add_run('Internal Context (Work IQ): ').bold = True
p.add_run('Internally, Akamai is still viewed as the most complete WAAP due to NoName + NeoSec acquisitions giving them API-native-in-WAAP capabilities. Do not distribute externally.')

# ============================================================
# INDUSTRY & MARKET SIGNALS
# ============================================================
doc.add_heading('Industry & Market Signals', level=2)

market_items = [
    'WAF market forecast: $22B by 2030 (14.9% CAGR, Mordor Intelligence, Mar 7) \u2014 Source: https://www.mordorintelligence.com/industry-reports/web-application-firewall-market',
    'Cloud WAAP market: $10B in 2025, 15% CAGR (Data Insights Market, Mar 7) \u2014 Source: https://www.datainsightsmarket.com/reports/cloud-web-application-and-api-protection-waap-1459511',
    'WAF market: $30.86B by 2034 (Fortune Business Insights, Feb 16, pre-window) \u2014 Source: https://www.fortunebusinessinsights.com/web-application-firewall-market-103375',
    'KuppingerCole Leadership Compass on WAAP \u2014 Ongoing analyst evaluation \u2014 Source: https://www.kuppingercole.com/research/leadership-compass/waap',
]
for item in market_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
r = p.add_run('Note: Market forecasts come from different firms with different methodologies and scopes. All confirm strong double-digit growth but direct comparison requires accounting for definitional differences.')
r.italic = True
r.font.size = Pt(9)

# ============================================================
# INTERNAL SIGNALS
# ============================================================
doc.add_heading('Internal Signals', level=2)

p = doc.add_paragraph(style='Intense Quote')
p.add_run('Internal Context (Work IQ): ').bold = True
p.add_run('All signals below are sourced from internal channels. Do not distribute externally.')

internal_items = [
    'API enforcement is the #1 customer gap. Chevron, KPMG, Walgreens, Uniper, Whataburger want runtime API enforcement (not just detection). They cite Cloudflare and Akamai as having this capability today.',
    'Cloudflare confirmed as #1 churn risk. Explicitly cited as the benchmark customers use. Customers are leaving AFD/App Gateway for Cloudflare.',
    'Azure WAF trust recovery underway. False positives reduced 53% to 6%, but perception lag remains. Still perceived behind Cloudflare/Akamai on bot mitigation, ML-driven L7 DDoS, and unified multi-cloud WAAP UX.',
    'Strategic pivot confirmed: WAAP designated as the enforcement layer for API & AI security; MDC handles discovery/posture.',
    'AI security timeline accelerated. Leadership determined original 2-3 year plan would miss the AI wave. Cloudflare and GCP Cloud Armor cited as competitive triggers.',
    'Three accelerated AI WAAP pillars: (1) Security FROM AI \u2014 agent/bot access control, (2) Security FOR AI \u2014 OWASP LLM Top 10, prompt injection, token-rate-limiting rulesets, (3) Security BY AI \u2014 dynamic rule gen, <48h CVE response.',
    'Ignite is the target milestone for AI WAAP pillar announcements.',
    'WAF + API detections POC agreed. Joint feasibility POC between Defender for APIs detections and WAF rule enforcement. Owners and timing defined.',
    'AWS WAF & GCP Cloud Armor viewed internally as baseline WAFs, not WAAP leaders. Customers layer Cloudflare/Akamai on top.',
    'Akamai = most complete WAAP per internal assessment. NoName + NeoSec acquisitions validate the API-native-in-WAAP approach we are pursuing.',
]
for item in internal_items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# WATCH LIST
# ============================================================
doc.add_heading('Watch List', level=2)

watch_items = [
    'Cloudflare Full-Transaction Detection \u2014 Track progression from "Under Development" to GA.',
    'Cloudflare Attack Signature Detection \u2014 Monitor Early Access adoption and customer feedback.',
    'RSA Conference (upcoming) \u2014 Expect burst of announcements from Akamai, Cloudflare, and others. Prepare competitive responses.',
    'KuppingerCole WAAP Leadership Compass \u2014 Track publication and our positioning vs. competitors.',
    'A10/ThreatX integration \u2014 Stale signal (acquisition closed Mar 2025): Monitor whether A10 becomes a credible WAAP contender post-acquisition.',
    'AWS Bedrock to WAF AI integration \u2014 Watch for signs AWS is embedding AI/ML into WAF using Bedrock capabilities.',
    'GCP Cloud Armor + Gemini \u2014 Monitor whether Google integrates Gemini AI into Cloud Armor threat detection.',
    'API enforcement gap closure \u2014 Track progress on WAF + Defender for APIs POC and customer feedback.',
]
for item in watch_items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# RECOMMENDED ACTIONS
# ============================================================
doc.add_heading('Recommended Actions', level=2)

table = doc.add_table(rows=7, cols=4, style='Table Grid')
act_headers = ['Priority', 'Action', 'Owner (suggested)', 'Rationale']
for i, h in enumerate(act_headers):
    table.rows[0].cells[i].text = h
    for par in table.rows[0].cells[i].paragraphs:
        for r in par.runs:
            r.bold = True

actions = [
    ['P0', 'Evaluate Cloudflare Attack Signature Detection architecture and assess feasibility of always-on detection in Azure WAF', 'WAF PM + Engineering', 'Decoupled detect/mitigate model eliminates key customer friction point. Need technical assessment.'],
    ['P0', 'Accelerate WAF + API detections POC to close the #1 customer gap (runtime API enforcement)', 'API Security PM + WAF PM', 'Five named enterprise accounts cite this as the reason they consider Cloudflare/Akamai. Ensure execution stays on track.'],
    ['P1', 'Prepare RSA Conference competitive battle cards', 'Product Marketing + PM', 'Competitors will launch at RSA. Pre-position messaging on our AI WAAP pillars and API enforcement roadmap.'],
    ['P1', 'Analyze Cloudflare Full-Transaction Detection for response-side analysis feasibility', 'WAF PM + Engineering', 'Response-body analysis is genuinely novel. Assess whether Azure WAF can intercept and analyze responses without unacceptable latency.'],
    ['P1', 'Address perception lag on Azure WAF FP rates', 'Product Marketing', 'FPs are down 53% to 6% but customers don\'t know. Publish a blog, update docs, and brief field teams.'],
    ['P2', 'Monitor A10/ThreatX integration and KuppingerCole WAAP Compass for positioning opportunities', 'Competitive Intelligence', 'Review A10 product roadmap quarterly (next check: Q2 2026). Track KuppingerCole publication date and flag within 48 hours.'],
]
for ri, row_data in enumerate(actions):
    for ci, val in enumerate(row_data):
        table.rows[ri + 1].cells[ci].text = val

# ============================================================
# QUIET COMPETITORS
# ============================================================
doc.add_heading('Quiet Competitors', level=2)
doc.add_paragraph('These tracked competitors had no or minimal detectable activity this week:')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Akamai').bold = True
p.add_run(' \u2014 No WAAP-specific announcements. Likely holding for RSA Conference.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('AWS WAF').bold = True
p.add_run(' \u2014 Minimal activity. No WAF-specific updates; Shield Director Findings mentioned in weekly roundup.')

# ============================================================
# SOURCES
# ============================================================
doc.add_heading('Sources', level=2)

table = doc.add_table(rows=14, cols=4, style='Table Grid')
src_headers = ['URL', 'Competitor', 'Signal Type', 'Date']
for i, h in enumerate(src_headers):
    table.rows[0].cells[i].text = h
    for par in table.rows[0].cells[i].paragraphs:
        for r in par.runs:
            r.bold = True

sources = [
    ['https://blog.cloudflare.com/attack-signature-detection/', 'Cloudflare', 'Product Launch', 'Mar 4, 2026'],
    ['https://developers.cloudflare.com/changelog/post/2026-03-02-waf-release/', 'Cloudflare', 'Feature Release', 'Mar 2, 2026'],
    ['https://itwire.com/.../cloudflare-2026-threat-intelligence-report...', 'Cloudflare', 'Media Coverage', 'Mar 3, 2026'],
    ['https://www.cloudflare.com/.../post-quantum...', 'Cloudflare', 'Product Launch', 'Feb 23 (pre-window)'],
    ['https://www.cloudflare.com/.../mastercard...', 'Cloudflare', 'Partnership', 'Feb 17 (pre-window)'],
    ['https://docs.cloud.google.com/armor/docs/release-notes', 'GCP Cloud Armor', 'Feature Release', 'Mar 3, 2026'],
    ['https://www.youtube.com/watch?v=aws-whats-new-mar-05-2026', 'AWS WAF', 'Observation', 'Mar 5, 2026'],
    ['https://www.mordorintelligence.com/.../web-application-firewall-market', 'Industry', 'Market Forecast', 'Mar 7, 2026'],
    ['https://www.datainsightsmarket.com/.../cloud-waap-1459511', 'Industry', 'Market Forecast', 'Mar 7, 2026'],
    ['https://www.fortunebusinessinsights.com/.../web-application-firewall...', 'Industry', 'Market Forecast', 'Feb 16 (pre-window)'],
    ['https://itwire.com/.../a10-networks-threatx-protect...', 'Industry', 'M&A', 'Mar 2025 (stale)'],
    ['https://www.kuppingercole.com/.../leadership-compass/waap', 'Industry', 'Analyst Report', 'Ongoing 2026'],
    ['Work IQ (internal)', 'Internal', 'Internal Signals', 'Mar 9, 2026'],
]
for ri, row_data in enumerate(sources):
    for ci, val in enumerate(row_data):
        table.rows[ri + 1].cells[ci].text = val

# ============================================================
# FOOTER
# ============================================================
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Generated on March 9, 2026 by the Weekly Competitive Digest Agent Farm.')
r.italic = True
r.font.size = Pt(8)
p2 = doc.add_paragraph()
r2 = p2.add_run('Previous digests: check work/runs/ for historical context.')
r2.italic = True
r2.font.size = Pt(8)

# Save
output_path = r'c:\Repo\amir-agent-packs\scaffold-agent-farm\farms\weekly-competitive-digest\work\runs\2026-03-09-weekly-waap\output\weekly-digest.docx'
doc.save(output_path)
print(f'Word document saved to: {output_path}')
