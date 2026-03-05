"""
Generate a polished DOCX competitive brief from the markdown source.
Uses python-docx to create a professionally styled document.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Paths ──────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "competitive-brief.docx")

# ── Colour palette ─────────────────────────────────────────────────────
MSFT_BLUE   = RGBColor(0x00, 0x78, 0xD4)
DARK_GREY   = RGBColor(0x33, 0x33, 0x33)
MED_GREY    = RGBColor(0x66, 0x66, 0x66)
LIGHT_GREY  = RGBColor(0xF2, 0xF2, 0xF2)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GREEN       = RGBColor(0x10, 0x7C, 0x10)
RED         = RGBColor(0xD1, 0x34, 0x38)
AMBER       = RGBColor(0xCA, 0x83, 0x00)
INTERNAL_BG = "FFF8E1"

# ── Helpers ────────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right with dict of sz, color, val."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="0078D4",
                     font_size=8, bold_first_col=False, alt_row_shading=True):
    """Create a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.color.rgb = DARK_GREY
            if bold_first_col and c_idx == 0:
                run.bold = True
            if alt_row_shading and r_idx % 2 == 1:
                set_cell_shading(cell, "F5F5F5")

    # Apply col widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)

    # Reduce cell margins
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)

    return table


def add_heading(doc, text, level=1, color=None):
    """Add a styled heading."""
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, color=None, size=10, space_after=6):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, bold_prefix="", level=0, size=9):
    """Add a bullet point with optional bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(size)
        run = p.add_run(text)
        run.font.size = Pt(size)
    else:
        run = p.add_run(text)
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_callout(doc, text, bg_color=INTERNAL_BG):
    """Add a callout/warning box using a single-cell table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = MED_GREY
    set_cell_shading(cell, bg_color)
    for edge in ["top", "bottom", "left", "right"]:
        kwargs = {edge: {"sz": "4", "color": "E0C060", "val": "single"}}
        set_cell_border(cell, **kwargs)
    doc.add_paragraph()  # spacer


def add_battle_card_table(doc, title, rows):
    """Add a battle card table with colour coding."""
    add_heading(doc, title, level=3, color=MSFT_BLUE)
    headers = ["Dimension", "We Win Because", "They Win Because", "Engineering Response"]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = WHITE
        set_cell_shading(cell, "0078D4")

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            if c_idx == 1 and val != "\u2014":
                run.font.color.rgb = GREEN
            elif c_idx == 2 and val != "\u2014":
                run.font.color.rgb = RED
            else:
                run.font.color.rgb = DARK_GREY
            if c_idx == 0:
                run.bold = True
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    # Col widths
    for row in table.rows:
        row.cells[0].width = Inches(1.2)
        row.cells[1].width = Inches(2.2)
        row.cells[2].width = Inches(2.2)
        row.cells[3].width = Inches(3.4)

    doc.add_paragraph()  # spacer


# ══════════════════════════════════════════════════════════════════════
# MAIN DOCUMENT BUILD
# ══════════════════════════════════════════════════════════════════════

doc = Document()

# ── Page setup ─────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width = Inches(11)
section.page_height = Inches(8.5)
section.orientation = WD_ORIENT.LANDSCAPE
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

# ── Default font ───────────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Segoe UI"
font.size = Pt(10)
font.color.rgb = DARK_GREY

# Heading styles
for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Segoe UI"
    hs.font.color.rgb = MSFT_BLUE if level <= 2 else DARK_GREY

# ══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Competitive Analysis Brief")
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = MSFT_BLUE

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Microsegmentation / Zero Trust Segmentation")
run.font.size = Pt(18)
run.font.color.rgb = MED_GREY

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for label, value in [
    ("Date: ", "March 2, 2026"),
    ("  |  Audience: ", "Engineering"),
    ("  |  Classification: ", "Internal"),
]:
    run = meta.add_run(label)
    run.font.size = Pt(11)
    run.font.color.rgb = MED_GREY
    run = meta.add_run(value)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GREY

doc.add_paragraph()

competitors_line = doc.add_paragraph()
competitors_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = competitors_line.add_run("Illumio  ·  Akamai Guardicore  ·  Zero Networks  ·  Microsoft Azure ZTS")
run.font.size = Pt(13)
run.font.color.rgb = MSFT_BLUE

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# HOW TO USE THIS BRIEF
# ══════════════════════════════════════════════════════════════════════

add_heading(doc, "How to Use This Brief", level=1)
add_styled_table(doc,
    headers=["Audience", "Focus Sections"],
    rows=[
        ["Engineering", "Comparison Matrix → Competitor Deep-Dives → Head-to-Head Analysis → Technical features"],
        ["Leadership", "Executive Summary → Strategic Recommendations → Gaps & Open Questions"],
        ["PM Peers", "Comparison Matrix → Competitor Deep-Dives → Roadmap Leverage"],
        ["Sales / GTM", "Battle Cards → Where We Win / Lose → Objection Handlers"],
    ],
    col_widths=[1.8, 7.2],
    font_size=9,
    bold_first_col=True,
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════

add_heading(doc, "Executive Summary", level=1)

add_para(doc,
    "The microsegmentation market has four significant players: Illumio (market leader, agent-based, "
    "best visibility), Akamai Guardicore (co-leader, process-level granularity, broadest platform "
    "coverage), Zero Networks (disruptor, agentless, MFA-gated, automated policies), and Microsoft "
    "Azure ZTS (late entrant, Azure-native, agentless, in Private Preview).",
    size=10)

add_para(doc,
    "Illumio and Guardicore dominate with proven enterprise scale and 5-10 year track records, "
    "but both require per-workload agents that add deployment and operational overhead. Zero Networks "
    "stands out through fully automated policy generation and unique MFA-gated lateral movement "
    "prevention — but lacks scale validation and process-level visibility.",
    size=10)

add_para(doc,
    "Microsoft Azure ZTS differentiates through zero-infrastructure deployment and native Azure "
    "enforcement at multiple layers (NSG + AVNM + Azure Firewall + NSP + Defender + Entra ID), "
    "but ships at GA with subnet-level granularity only, no PaaS/AKS segmentation, and Azure-only scope.",
    size=10)

p = doc.add_paragraph()
run = p.add_run("Our strongest play: ")
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = GREEN
run = p.add_run("Azure-centric organizations that want microsegmentation without another vendor, agent, or console.")
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run("Our biggest risk: ")
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RED
run = p.add_run("Feature gap perception at GA — competitors offer process-level granularity, multi-cloud, and identity-aware segmentation today.")
run.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# MARKET LANDSCAPE
# ══════════════════════════════════════════════════════════════════════

add_heading(doc, "Market Landscape", level=1)

add_heading(doc, "Market Context", level=2)
add_bullet(doc, "Microsegmentation is a core Zero Trust pillar, accelerated by ransomware proliferation and federal mandates (US EO 14028).")
add_bullet(doc, "Gen 1 (2013+): ", bold_prefix="")
p = doc.add_paragraph(style="List Bullet")
run = p.add_run("Gen 1 (2013+): ")
run.bold = True
run.font.size = Pt(9)
run = p.add_run("Agent-based, policy-driven (Illumio, Guardicore)")
run.font.size = Pt(9)
p.paragraph_format.space_after = Pt(3)

p = doc.add_paragraph(style="List Bullet")
run = p.add_run("Gen 2 (2019+): ")
run.bold = True
run.font.size = Pt(9)
run = p.add_run("Agentless, identity-aware, automated (Zero Networks)")
run.font.size = Pt(9)
p.paragraph_format.space_after = Pt(3)

p = doc.add_paragraph(style="List Bullet")
run = p.add_run("Gen 3 (2025+): ")
run.bold = True
run.font.size = Pt(9)
run = p.add_run("Cloud-native, platform-integrated (Microsoft Azure ZTS)")
run.font.size = Pt(9)
p.paragraph_format.space_after = Pt(3)

add_bullet(doc, "Identity + network convergence is the major architectural trend. Zero Networks pioneered this; Microsoft's Horizon 3 roadmap targets it.")

add_heading(doc, "Market Positioning", level=2)
add_styled_table(doc,
    headers=["Vendor", "Est. Mindshare", "Positioning", "Primary Buyer"],
    rows=[
        ["Illumio", "~29%", 'Market leader, "breach containment platform"', "Large enterprise CISO teams"],
        ["Guardicore (Akamai)", "Co-leader", "Process-level segmentation, broadest coverage", "Enterprise security / compliance teams"],
        ["Zero Networks", "~4% (growing)", "Identity-first innovator, automated", "Lean security teams, mid-market to upper mid-market"],
        ["Microsoft Azure ZTS", "Late entrant", "Azure-native, zero friction", "Azure-centric cloud / network teams"],
    ],
    font_size=9,
    bold_first_col=True,
)
add_callout(doc, "⚠ Mindshare percentages are from internal intelligence and have not been validated against published analyst data.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# COMPARISON MATRIX (large table — landscape works well here)
# ══════════════════════════════════════════════════════════════════════

add_heading(doc, "Comparison Matrix", level=1)

cmp_headers = ["Dimension", "Microsoft Azure ZTS", "Illumio", "Akamai Guardicore", "Zero Networks"]
cmp_rows = [
    ["Architecture",
     "Azure PaaS (AVNM add-on). No agents, no customer infra. NSG enforcement (GA).",
     "Host agent (VEN) on every workload. PCE controller.",
     "Host agent on every workload. Centralized mgmt server.",
     "Agentless. SaaS/VA. Remote OS firewalls (WMI/WinRM)."],
    ["Enforcement Granularity",
     "Subnet-level (GA). Workload-level post-GA. No process visibility.",
     "Process-level. VEN programs OS firewall per-process/port.",
     "Process-level. Process + binary hash identification.",
     "IP:port level. No process visibility. Default-deny."],
    ["Policy Authoring",
     "AI auto-seg + manual Portal. Intent-based. ARM/Bicep/Terraform.",
     "Semi-automated. Illumination → label-based → staged rollout.",
     "Semi-automated. Reveal → label-based → staged rollout.",
     "Fully automated ML. Human approves/overrides only."],
    ["Deployment Time",
     "Minutes (enable, onboard VNets).",
     "Weeks–months (agents + 2-4 wk discovery).",
     "Weeks–months (similar to Illumio).",
     "Hours–days (agentless, auto-learning)."],
    ["Multi-Environment",
     "Azure only (GA). Multi-cloud H3 (CY 2028+).",
     "On-prem, Azure, AWS, GCP, containers.",
     "On-prem, Azure, AWS, GCP, containers, IoT/OT.",
     "On-prem, Azure, AWS, GCP."],
    ["Identity Integration",
     "None at GA. Entra ID in H3.",
     "None native.",
     "None native.",
     "MFA-gated access (unique). AD/Azure AD aware."],
    ["Visibility/Mapping",
     "Communication Graph (flow-log, IP:port). Portal dashboards.",
     "Illumination Map (best-in-class). Process-level interactive.",
     "Reveal Map (strong). Process + binary hash graph.",
     "Traffic analytics/reporting. Automation-first."],
    ["PaaS Segmentation",
     "H2 (NSP for Storage/SQL/KV). Not at GA.",
     "Via PE + Azure Firewall integration. Manual.",
     "Via PE + agents. Manual.",
     "Via NSG/firewall for PE services. Auto."],
    ["AKS / Kubernetes",
     "H2 (namespace/pod). Not at GA.",
     "Kubelink.",
     "DaemonSet (pod-level). Strongest K8s.",
     "Limited. No specialized K8s."],
    ["Breach Simulation",
     "None.",
     "None.",
     "Infection Monkey (OSS BAS). Unique.",
     "None."],
    ["Anomaly Detection",
     "H2 (inter-segment anomaly). Not at GA.",
     "AI Security Graph (new, limited).",
     "Behavioral analysis + deception honeypots.",
     "Default-deny + MFA = block, not detect."],
    ["Pricing (100 wklds/yr)",
     "TBD (expected lowest).",
     "~$35K–$50K",
     "~$16.8K",
     "~$20K"],
    ["Pricing (500 wklds/yr)",
     "TBD",
     "~$175K–$250K",
     "~$84K",
     "~$100K"],
    ["Scale Track Record",
     "~25 preview customers.",
     "200K+ workloads. Fortune 100.",
     "Large enterprise. Akamai ($3.8B).",
     "Mid/upper-mid market. No 100K+ refs."],
    ["IaC / DevOps",
     "ARM, Bicep, Terraform, Azure Policy, GitHub Actions.",
     "Terraform, REST API, CI/CD.",
     "REST API, some IaC.",
     "REST API. Less IaC maturity."],
]

add_styled_table(doc, cmp_headers, cmp_rows,
                 col_widths=[1.6, 2.1, 2.1, 2.1, 2.1],
                 font_size=7, bold_first_col=True)

add_callout(doc, "⚠ Pricing caveat: All pricing figures are internal estimates. Validate before customer-facing use.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# COMPETITOR DEEP-DIVES
# ══════════════════════════════════════════════════════════════════════

add_heading(doc, "Competitor Deep-Dives", level=1)

# ── Illumio ────────────────────────────────────────────────────────────
add_heading(doc, "Illumio", level=2)
add_para(doc, "Architecture: Agent-based (VEN on every workload). Programs host OS firewall (WFP/iptables). Centralized PCE. Process-level telemetry and enforcement.", bold=True, size=9)

add_para(doc, "Key Technical Strengths:", bold=True, size=10, color=GREEN)
add_bullet(doc, " Industry gold standard for traffic visualization. Real-time interactive dependency graph. Process-level.", bold_prefix="Illumination Map:")
add_bullet(doc, " 4-dimensional (Role/App/Env/Location). IP-decoupled, human-readable policies.", bold_prefix="Label taxonomy:")
add_bullet(doc, ' "What-if" engine for safe rollout. Models rule changes before enforcement.', bold_prefix="Policy simulation:")
add_bullet(doc, " 200K+ workloads proven. PCE scales horizontally. Fail-closed agents.", bold_prefix="Scale:")
add_bullet(doc, " Forrester TEI claims 111% over 3 years. (Vendor-commissioned.)", bold_prefix="ROI:")

add_para(doc, "Key Technical Weaknesses:", bold=True, size=10, color=RED)
add_bullet(doc, "Agent lifecycle at scale (VEN compatibility per OS update, ephemeral workload orchestration).")
add_bullet(doc, "2-4 week mandatory discovery phase.")
add_bullet(doc, "No identity/MFA integration. Credential theft within allowed segments undetected.")
add_bullet(doc, 'Azure Firewall integration: ~25 customers after 2+ years. Adoption "modest."')

add_callout(doc, 'Engineering Takeaway: Illumio\'s Illumination map and process-level granularity are what engineers will compare Azure ZTS against. The question they\'ll ask: "Can Azure ZTS show me which process is communicating?" Answer: no.', "E8F0FE")

# ── Guardicore ─────────────────────────────────────────────────────────
add_heading(doc, "Akamai Guardicore", level=2)
add_para(doc, "Architecture: Agent-based. Process + binary hash identification. Broadest OS/platform matrix: Windows 2003+, CentOS 5+, K8s, IoT/OT.", bold=True, size=9)

add_para(doc, "Key Technical Strengths:", bold=True, size=10, color=GREEN)
add_bullet(doc, " Forensic-grade visibility. Identifies exact binary behind each flow.", bold_prefix="Process + binary hash:")
add_bullet(doc, " Open-source BAS tool. Tests segmentation by simulating real attacks. Unique.", bold_prefix="Infection Monkey:")
add_bullet(doc, " Network-level enforcement for unmanaged devices. Critical for manufacturing, healthcare.", bold_prefix="IoT/OT:")
add_bullet(doc, " Built-in honeypots detect lateral movement.", bold_prefix="Deception:")
add_bullet(doc, " $3.8B public company. No vendor risk.", bold_prefix="Akamai backing:")

add_para(doc, "Key Technical Weaknesses:", bold=True, size=10, color=RED)
add_bullet(doc, "Same agent overhead as Illumio (weeks–months deployment).")
add_bullet(doc, "Post-acquisition integration uncertainty.")
add_bullet(doc, "No identity/MFA.")

add_callout(doc, "Engineering Takeaway: Guardicore's binary-hash identification and Infection Monkey are capabilities no other competitor can match. For PCI DSS / HIPAA scope reduction, the forensic visibility is decisive.", "E8F0FE")

doc.add_page_break()

# ── Zero Networks ──────────────────────────────────────────────────────
add_heading(doc, "Zero Networks", level=2)
add_para(doc, "Architecture: Agentless. SaaS controller or VA. Orchestrates remote OS firewalls (WMI/WinRM, SSH) and network device APIs (Palo Alto DAGs). Zero host performance impact.", bold=True, size=9)

add_para(doc, "Key Technical Strengths:", bold=True, size=10, color=GREEN)
add_bullet(doc, " Out-of-baseline connections trigger MFA challenge. Blocks credential-theft lateral movement.", bold_prefix="MFA-gated access:")
add_bullet(doc, " Traffic learning → auto-generated allow-lists → auto-enforcement. Near-zero manual work.", bold_prefix="Automated ML policy:")
add_bullet(doc, " RDP/SSH/admin ports closed until JIT MFA-verified. Strongest default posture.", bold_prefix="Default-deny all ports:")
add_bullet(doc, " Hours–days (vendor claims '15 minutes' — real-world may vary).", bold_prefix="Deployment speed:")
add_bullet(doc, " Segmentation + identity segmentation + ZTNA in one product.", bold_prefix="Unified platform:")

add_para(doc, "Key Technical Weaknesses:", bold=True, size=10, color=RED)
add_bullet(doc, "IP:port enforcement only. No process visibility.")
add_bullet(doc, "WMI/WinRM dependency. Hardened environments may restrict these channels.")
add_bullet(doc, "No 100K+ workload references. OS firewall rule-count scaling limits.")
add_bullet(doc, "Limited interactive visualization. Automation-first design philosophy.")
add_bullet(doc, "No native Azure integration. No Marketplace presence.")

add_callout(doc, "Engineering Takeaway: Zero Networks sets the automation and UX bar. Their MFA-gated access addresses the credential-theft blind spot in every other competitor. The tradeoff: no process-level granularity and no interactive traffic maps.", "E8F0FE")

# ── Microsoft Azure ZTS ───────────────────────────────────────────────
add_heading(doc, "Microsoft Azure ZTS (Our Product)", level=2)
add_para(doc, "Architecture: Azure PaaS. AVNM add-on. Billable Networking SKU. Agentless. NSG enforcement (GA). NSG flow logs → Traffic Analytics → Communication Graph → AI auto-segmentation with explainability.", bold=True, size=9)

add_para(doc, "GA Scope (Horizon 1):", bold=True, size=10, color=MSFT_BLUE)
ga_rows = [
    ["Subscription/VNet onboarding via AVNM", "GA"],
    ["VM + internet-service inventory", "GA"],
    ["Complete Communication Graph (historical + continuous)", "GA"],
    ["AI-driven auto-segmentation with explainability", "GA"],
    ["Workload labeling (Env / App / Role)", "GA"],
    ["Intent-based policy → NSG rules", "GA"],
    ["Subnet-level enforcement", "GA"],
    ["RBAC, logging, metering, ARM APIs", "GA"],
    ["Azure Portal experience", "GA"],
]
add_styled_table(doc, ["Capability", "Status"], ga_rows,
                 col_widths=[6.0, 1.5], font_size=8, bold_first_col=True, header_color="107C10")

doc.add_paragraph()
add_para(doc, "Post-GA Roadmap:", bold=True, size=10, color=MSFT_BLUE)
add_styled_table(doc,
    headers=["Horizon", "Timeline", "Key Capabilities"],
    rows=[
        ["H2", "CY 2027+", "PaaS segmentation (NSP), AKS (namespace/pod), Azure Firewall L3-L7, AI policy recs, anomaly detection"],
        ["H3", "CY 2028+", "Multi-cloud (AWS/GCP), identity-aware (Entra ID), on-prem (MDE agents)"],
    ],
    col_widths=[1.0, 1.5, 6.5], font_size=9, bold_first_col=True)

doc.add_paragraph()
add_para(doc, "ML/AI Timeline:", bold=True, size=10, color=MSFT_BLUE)
add_styled_table(doc,
    headers=["Milestone", "Timeline"],
    rows=[
        ["ML segment discovery pipeline", "Private Preview (now)"],
        ["Explainability + enhanced microseg", "Public Preview Q3 CY 2026"],
        ["MCP server for AI workflows", "1H CY 2027"],
        ["ML policy recommendations", "2H CY 2027"],
        ["Anomaly detection", "2H CY 2027"],
    ],
    col_widths=[5.0, 4.0], font_size=9, bold_first_col=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# HEAD-TO-HEAD ANALYSIS
# ══════════════════════════════════════════════════════════════════════

add_heading(doc, "Head-to-Head Analysis", level=1)

add_heading(doc, "Where We Win", level=2, color=GREEN)
add_styled_table(doc,
    headers=["Scenario", "Why Azure ZTS Wins"],
    rows=[
        ["Azure-only IaaS estates", "Zero deployment friction. No agents. Lower TCO. Native Portal."],
        ["Large Azure subscriptions", "Hyperscale governance. AVNM distributes policies at any scale."],
        ["DevOps / IaC-first teams", "ARM, Bicep, Terraform, Azure Policy. Segmentation-as-code."],
        ["Teams already using AVNM", "Natural extension. Same management plane."],
        ["Cost-conscious buyers", "No agent license. No infra. Azure consumption pricing."],
        ["Compliance (Azure-only)", "Azure compliance posture. ARM audit logs. Azure Policy enforcement."],
    ],
    col_widths=[2.5, 6.5], font_size=9, bold_first_col=True, header_color="107C10")

doc.add_paragraph()

add_heading(doc, "Where We Lose", level=2, color=RED)
add_styled_table(doc,
    headers=["Scenario", "Winner", "Why"],
    rows=[
        ["Process-level visibility", "Illumio / Guardicore", "Agents identify communicating processes. We see IP:port only."],
        ["Hybrid / multi-cloud", "Illumio / Guardicore", "One policy across on-prem + multi-cloud. We're Azure-only."],
        ["Credential theft mitigation", "Zero Networks", "MFA-gated access. No identity integration at GA."],
        ["IoT/OT segmentation", "Guardicore", "Network enforcement for unmanaged devices. We don't cover this."],
        ["Interactive flow visualization", "Illumio / Guardicore", "Process-level interactive maps vs. our IP:port flow-log CCG."],
        ["Kubernetes workloads", "Guardicore / Illumio", "Pod-level (DaemonSet / Kubelink). No AKS support until H2."],
        ["Breach simulation", "Guardicore", "Infection Monkey BAS. No equivalent from us."],
        ["Zero manual policy work", "Zero Networks", "Fully automated ML. Our AI is semi-automated."],
    ],
    col_widths=[2.2, 2.0, 4.8], font_size=9, bold_first_col=True, header_color="D13438")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# INTERNAL CONTEXT
# ══════════════════════════════════════════════════════════════════════

add_heading(doc, "Internal Context", level=1)
add_callout(doc, "⚠ INTERNAL — Do not distribute externally. Contains Work IQ intelligence.", "FFF3E0")

add_para(doc, "Our Differentiators:", bold=True, size=10, color=GREEN)
add_bullet(doc, "Native enforcement at multiple Azure layers (NSG + AVNM + Firewall + NSP + Policy + Defender + Entra)")
add_bullet(doc, "Platform context advantage (Azure knows resources, tags, topology, identity, compliance)")
add_bullet(doc, "Hyperscale governance (no customer-managed controllers or agents)")
add_bullet(doc, "Lower TCO for Azure-centric customers (internal estimate — build defensible TCO model)")

add_para(doc, "Competitive Concerns:", bold=True, size=10, color=RED)
add_bullet(doc, "Late market entry (5-10 year competitor head start)")
add_bullet(doc, "Feature gap at GA (subnet-level, no PaaS/AKS, no identity, no Firewall enforcement)")
add_bullet(doc, "Azure-only scope in a hybrid/multi-cloud world")
add_bullet(doc, "Positioning confusion (NSG vs ASG vs AVNM vs Firewall vs NSP vs ZTS)")

add_para(doc, "Customer Intelligence:", bold=True, size=10, color=AMBER)
add_bullet(doc, "Bridgewater Associates drove Illumio-Azure Firewall integration; potential ZTS design partner")
add_bullet(doc, '~25 Private Preview customers; adoption "modest"')
add_bullet(doc, "No documented deal losses to competitors")

add_para(doc, "Roadmap Leverage:", bold=True, size=10, color=MSFT_BLUE)
add_bullet(doc, 'H1 (GA): "See, Understand & Enforce" for Azure IaaS')
add_bullet(doc, "H2: Closes major gaps (PaaS, AKS, Firewall, AI recs) vs. Illumio/Guardicore")
add_bullet(doc, "H3: Addresses two biggest objections (multi-cloud, identity-aware)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# BATTLE CARDS
# ══════════════════════════════════════════════════════════════════════

add_heading(doc, "Battle Cards", level=1)

add_battle_card_table(doc, "vs. Illumio", [
    ["Deployment", "Zero agents, zero infra. Portal in minutes.", "Requires agent rollout (weeks–months).",
     '"Azure ZTS eliminates weeks of agent deployment and lifecycle management."'],
    ["TCO", "No agent license, no PCE infra, no patching.", "Most expensive (~$35-50K/100 wklds).",
     '"For 500 Azure workloads, Illumio costs $175-250K/yr + infra. Azure ZTS eliminates that."'],
    ["Visibility", "\u2014", "Illumination Map: process-level, interactive, real-time.",
     '"Our CCG covers traffic topology from flow logs. For process-level, pair with MDE host telemetry."'],
    ["Granularity", "\u2014", "Process-level enforcement via host agents.",
     '"Subnet-level covers common IaaS scenarios. Workload-level on roadmap. MDE for host-level."'],
    ["Multi-cloud", "\u2014", "On-prem + Azure + AWS + GCP, one policy.",
     '"For Azure: deepest integration. Multi-cloud: H3 roadmap. Today: Azure ZTS + existing tooling."'],
    ["IaC", "ARM, Bicep, Terraform, Azure Policy. Native.", "Terraform provider, less deep.",
     '"Define segmentation in the same Bicep/Terraform as your infra."'],
    ["Scale", "Azure fabric. No controller sizing.", "200K+ workloads proven.",
     '"Azure enforcement is in the fabric. Scaling is automatic."'],
])

add_battle_card_table(doc, "vs. Akamai Guardicore", [
    ["Deployment", "Zero agents, zero infra.", "Agent on every workload.",
     '"Azure ZTS eliminates weeks of agent deployment and lifecycle management."'],
    ["TCO", "No agent license, no mgmt server.", "More affordable than Illumio (~$16.8K/100 wklds/yr).",
     '"Even at Guardicore pricing, you pay for agents + infra."'],
    ["Process visibility", "\u2014", "Process + binary hash. Forensic-grade.",
     '"Pair Azure ZTS network segmentation with MDE host telemetry for complementary coverage."'],
    ["BAS", "\u2014", "Infection Monkey (OSS). Unique.",
     '"Infection Monkey is OSS and vendor-neutral — works with any segmentation, including Azure ZTS."'],
    ["IoT/OT", "\u2014", "Network enforcement for unmanaged devices.",
     '"For IoT/OT: Defender for IoT is the Azure-native answer."'],
    ["Platform breadth", "\u2014", "Broadest: Windows, Linux, K8s, IoT/OT, legacy.",
     '"Azure ZTS for Azure. Complementary tool for non-Azure."'],
])

doc.add_page_break()

add_battle_card_table(doc, "vs. Zero Networks", [
    ["Azure integration", "Native: Portal, RBAC, ARM, Bicep, Terraform.", "Third-party overlay, separate console.",
     '"Azure ZTS is alongside your NSGs, Firewall, AVNM. No new console."'],
    ["Enforcement depth", "Multi-layer: NSG + AVNM + Firewall (H2) + NSP (H2).", "Single layer: remote OS firewalls.",
     '"Multiple Azure enforcement planes vs. one remote firewall."'],
    ["Visibility", "Communication Graph: traffic topology.", "Analytics/reporting only, minimal maps.",
     '"Our CCG shows full topology. Their philosophy: trust the ML."'],
    ["Identity / MFA", "\u2014", "MFA-gated access. Blocks credential theft.",
     '"Innovative. Not at GA. Interim: Entra CA + PIM. H3: native Entra integration."'],
    ["Automation", "AI semi-automated.", "Fully automated ML. Near-zero manual effort.",
     '"We offer explainability (confidence scores, reasoning). Their approach is less transparent."'],
    ["Scale / maturity", "Azure hyperscale. Microsoft backing.", "No 100K+ refs. Private company risk.",
     '"For critical security infra: Microsoft\'s backing and Azure\'s fabric."'],
    ["Mgmt channel risk", "NSG enforcement in fabric. No workload access.", "WMI/WinRM must be open.",
     '"No management channel to the workload needed."'],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# STRATEGIC RECOMMENDATIONS
# ══════════════════════════════════════════════════════════════════════

add_heading(doc, "Strategic Recommendations", level=1)

recs = [
    ("1. Close the visibility gap messaging proactively",
     'Illumio\'s Illumination map is the #1 feature security evaluators look for. Position the Communication '
     'Graph as "cloud-native traffic analytics at Azure scale" and investigate whether MDE process telemetry '
     'can be surfaced alongside ZTS data. (Validate MDE + ZTS integration feasibility with the product team.)'),
    ("2. Accelerate workload-level granularity",
     "Subnet-level enforcement at GA is the single biggest feature gap. Treat AVNM workload-level enforcement "
     "as critical-path. If it can't unblock before GA, communicate a clear timeline and document NSG-per-VM workarounds."),
    ("3. Lead with 'zero friction' for greenfield Azure buyers",
     "Target organizations with no microsegmentation today — the 'do nothing' segment. The win is greenfield "
     "adoption, not competitive displacement. The #1 barrier to microseg adoption is complexity; Azure ZTS directly addresses this."),
    ("4. Build 'Azure ZTS + Defender + Entra' reference architecture",
     "Before H3 delivers native identity-aware segmentation, document a defense-in-depth architecture: ZTS (network) + "
     "Defender (posture + response) + Entra (identity CA). Acknowledge honestly that this is a multi-product approach."),
    ("5. Clarify the Azure networking tool landscape",
     "Publish a decision tree: when to use NSG vs ASG vs AVNM vs Firewall vs NSP vs ZTS. Position ZTS as the "
     'policy intelligence layer — the "brain" that tells NSG/Firewall (the "muscles") what to enforce.'),
]

for title, body in recs:
    add_para(doc, title, bold=True, size=11, color=MSFT_BLUE, space_after=3)
    add_para(doc, body, size=10, space_after=10)

# ══════════════════════════════════════════════════════════════════════
# GAPS & OPEN QUESTIONS
# ══════════════════════════════════════════════════════════════════════

add_heading(doc, "Gaps & Open Questions", level=1)

gaps = [
    "Azure ZTS pricing — Not announced. Model vs. Guardicore (~$16.8K/100 wklds/yr).",
    "AVNM workload-level timeline — When does NIC-level enforcement unblock?",
    "MDE + ZTS integration — Can MDE data surface in ZTS Communication Graph?",
    "Zero Networks technical deep-dive — How does it manipulate Azure NSGs? IAM requirements? Scale?",
    "Guardicore Azure integration depth — Azure APIs or purely agent-based?",
    "Preview customer feedback — NPS, feature requests, friction points from ~25 customers.",
    "Illumio Azure Firewall adoption — Why only ~25 after 2+ years? Signal for Azure-native demand?",
    "Win/loss data — No competitive encounters documented. Need field intelligence.",
    'TCO model — "60-70% lower" claim needs rigorous validation.',
    'EMA survey verification — "82.8% automated policy" statistic is unverified.',
]

for gap in gaps:
    # use checkbox style bullet
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run("☐  " + gap)
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_GREY
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SOURCES
# ══════════════════════════════════════════════════════════════════════

add_heading(doc, "Sources", level=1)

add_styled_table(doc,
    headers=["Vendor", "Source", "Type"],
    rows=[
        ["Illumio", "illumio.com/products", "Public"],
        ["Illumio", "illumio.com/products/zero-trust-segmentation", "Public"],
        ["Guardicore", "akamai.com/products/akamai-guardicore-segmentation", "Public"],
        ["Zero Networks", "zeronetworks.com/products", "Public"],
        ["Zero Networks", "zeronetworks.com/platform", "Public"],
        ["Microsoft", "learn.microsoft.com — AVNM Security Admin Rules", "Public"],
        ["Microsoft", "learn.microsoft.com — Zero Trust Networking", "Public"],
        ["Internal", "Work IQ — competitive landscape, pricing, differentiators, threats, win/loss, roadmap", "Internal"],
        ["PM Resource", "work/resources/compete.md.txt", "LLM-generated (verified selectively)"],
    ],
    col_widths=[1.5, 5.5, 2.0], font_size=9, bold_first_col=True)

# ── Footer ─────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Generated on March 2, 2026  |  Microsoft Confidential — Internal Use Only")
run.font.size = Pt(8)
run.font.color.rgb = MED_GREY
run.italic = True

# ══════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════

doc.save(OUTPUT_PATH)
print(f"✓ Saved: {OUTPUT_PATH}")
