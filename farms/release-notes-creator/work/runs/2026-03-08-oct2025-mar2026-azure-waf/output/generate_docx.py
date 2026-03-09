"""Generate release-notes.docx from structured content."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


def add_title_page(doc: Document) -> None:
    """Add a centered title page."""
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Azure WAF Release Notes")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 0x78, 0xD4)  # Azure blue

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("October 2025 – March 2026")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_page_break()


def add_section_heading(doc: Document, text: str) -> None:
    """Add a styled section heading."""
    heading = doc.add_heading(text, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0x78, 0xD4)


def add_feature(doc: Document, title: str, body: str, note: str = "") -> None:
    """Add a feature entry with bold title, body paragraph, and optional italic note."""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)

    body_p = doc.add_paragraph(body)
    body_p.paragraph_format.space_after = Pt(4)

    if note:
        note_p = doc.add_paragraph()
        run = note_p.add_run(note)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_bullet_list(doc: Document, items: list[tuple[str, str]]) -> None:
    """Add a bulleted list with bold label and description."""
    for label, desc in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(desc)


def add_footer(doc: Document, text: str) -> None:
    """Add footer text to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def build_document() -> Document:
    """Build the full release notes document."""
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    add_title_page(doc)

    # --- New Features ---
    add_section_heading(doc, "New Features")

    add_feature(
        doc,
        "JavaScript Challenge for Bot Protection (GA)",
        "Application Gateway WAF now supports JavaScript Challenge as a generally "
        "available response action. When suspicious traffic is detected, the WAF "
        "issues a lightweight JavaScript challenge to the client browser, helping "
        "distinguish legitimate users from automated bots without disrupting the "
        "browsing experience.",
    )

    add_feature(
        doc,
        "Layer 7 DDoS Protection with Automatic Baseline Learning",
        "Azure WAF introduces new Layer 7 DDoS detection and mitigation capabilities:",
    )
    add_bullet_list(doc, [
        ("Detection enhancements",
         "Improved detection algorithms identify application-layer DDoS attacks "
         "that bypass caching layers to target origin servers, with greater accuracy "
         "for volumetric and sophisticated attack patterns."),
        ("Automatic baseline mitigation",
         "The system learns normal traffic patterns during steady-state periods and "
         "automatically triggers mitigation when anomalous traffic deviates from "
         "established baselines \u2014 no manual threshold tuning required."),
        ("Attack-signature mitigation",
         "A complementary real-time attack signature detection and response layer "
         "provides defense-in-depth against application-layer DDoS attacks."),
    ])

    add_feature(
        doc,
        "IPv6 Support for Application Gateway WAF (GA)",
        "Application Gateway WAF now fully supports IPv6 traffic at general "
        "availability. Customers with dual-stack or IPv6-only environments can "
        "apply the same WAF policy protections to IPv6 traffic that were previously "
        "available only for IPv4, enabling WAF protection across all network "
        "configurations.",
    )

    # --- Improvements ---
    add_section_heading(doc, "Improvements")

    add_feature(
        doc,
        "Default Rule Set (DRS) 2.2",
        "Azure WAF ships Default Rule Set version 2.2 for both Azure Front Door "
        "and Application Gateway. DRS 2.2 delivers updated detection rules with "
        "improved accuracy, reduced false positives, and expanded coverage for the "
        "latest web application attack vectors.",
    )

    add_feature(
        doc,
        "Core Rule Set (CRS) 4.0 Support",
        "Azure WAF adds support for OWASP Core Rule Set (CRS) 4.0, the latest "
        "major version of the OWASP ModSecurity Core Rule Set. CRS 4.0 delivers "
        "modernized rule definitions, improved categorization, and reduced false "
        "positives for common web frameworks.",
    )

    add_feature(
        doc,
        "Azure Monitor Workbooks for WAF (GA)",
        "Interactive Azure Monitor Workbooks for WAF are now generally available "
        "for both Azure Front Door and Application Gateway. These dashboards "
        "provide visibility into rule evaluations, block reasons, and traffic "
        "patterns. Aggregated false-positive pattern views help customers tune "
        "their WAF configurations with confidence.",
    )

    # --- Preview ---
    add_section_heading(doc, "Preview")

    add_feature(
        doc,
        "WAF Exception Lists (Allow Lists) for Application Gateway and Front Door",
        "Azure WAF introduces exception lists (also known as allow lists) for both "
        "Application Gateway and Azure Front Door. Customers can define fine-grained "
        "exclusions to exempt specific requests, fields, or URI patterns from WAF "
        "rule inspection \u2014 directly addressing a common source of false positives. "
        "For Azure Front Door, new resource provider APIs enable programmatic and "
        "template-based management of exception entries, supporting "
        "infrastructure-as-code workflows.",
    )

    add_feature(
        doc,
        "CAPTCHA Challenge for Bot Protection (Private Preview)",
        "Application Gateway WAF introduces CAPTCHA challenge as a new response "
        "action. Customers can configure WAF policies to present CAPTCHA challenges "
        "to suspicious requests, adding an interactive human-verification layer "
        "alongside the existing JavaScript Challenge. Together, these capabilities "
        "provide a layered bot defense strategy for web applications.",
        "Customers interested in early access should contact their Microsoft account team.",
    )

    add_feature(
        doc,
        "Rate Limiting with X-Forwarded-For (XFF) Header Support",
        "WAF rate limiting now supports the X-Forwarded-For header as a key for "
        "throttling decisions. This ensures accurate per-client rate limiting for "
        "traffic that passes through proxies or load balancers, using the true "
        "client IP address instead of the intermediary\u2019s address.",
    )

    # --- Known Limitations ---
    add_section_heading(doc, "Known Limitations (Preview Features)")

    doc.add_paragraph(
        "WAF Exception Lists (Allow Lists): The scope of supported match variables "
        "may be limited during preview. Customers should validate their exclusion "
        "rules against their specific traffic patterns.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "CAPTCHA Challenge: Browser compatibility scope has not been finalized. "
        "Customers participating in the private preview should test across their "
        "target browser matrix.",
        style="List Bullet",
    )

    # --- Documentation Updates ---
    add_section_heading(doc, "Documentation Updates")

    doc.add_paragraph(
        "Slow HTTP / Slowloris protection guidance: Improved documentation and "
        "guidance for configuring protection against Slow HTTP and Slowloris attacks "
        "is in progress. This addresses a recurring customer request for clearer "
        "best-practice recommendations.",
        style="List Bullet",
    )

    # Footer
    add_footer(doc, "Generated: March 8, 2026")

    return doc


if __name__ == "__main__":
    output_dir = Path(__file__).parent
    doc = build_document()
    output_path = output_dir / "release-notes.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")
